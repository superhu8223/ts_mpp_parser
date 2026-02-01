#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""读取docx文档内容并保存到文本文件"""

from docx import Document

def read_docx_to_file(docx_path, output_txt):
    """读取docx文档并保存到文本文件"""
    doc = Document(docx_path)
    
    with open(output_txt, 'w', encoding='utf-8') as f:
        f.write(f"{'='*60}\n")
        f.write(f"文档: {docx_path}\n")
        f.write(f"{'='*60}\n\n")
        
        for para in doc.paragraphs:
            if para.text.strip():
                f.write(para.text + "\n")
        
        # 读取表格内容
        if doc.tables:
            f.write("\n\n表格内容:\n")
            for i, table in enumerate(doc.tables):
                f.write(f"\n--- 表格 {i+1} ---\n")
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    f.write(" | ".join(row_data) + "\n")

if __name__ == "__main__":
    # 读取prompt_v2.docx
    read_docx_to_file(
        r"d:\wsp\test\ts_mpp_parser_vscode\prompt_v2.docx",
        r"d:\wsp\test\ts_mpp_parser_vscode\prompt_v2_extracted.txt"
    )
    print("已将 prompt_v2.docx 内容提取到 prompt_v2_extracted.txt")
    
    # 读取TsMpp开发参考.docx
    read_docx_to_file(
        r"d:\wsp\test\ts_mpp_parser_vscode\TsMpp开发参考.docx",
        r"d:\wsp\test\ts_mpp_parser_vscode\TsMpp_extracted.txt"
    )
    print("已将 TsMpp开发参考.docx 内容提取到 TsMpp_extracted.txt")
