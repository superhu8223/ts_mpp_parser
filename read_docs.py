#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""读取docx文档内容的临时脚本"""

from docx import Document

def read_docx(file_path):
    """读取docx文档并打印内容"""
    doc = Document(file_path)
    print(f"\n{'='*60}")
    print(f"文档: {file_path}")
    print(f"{'='*60}\n")
    
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    
    # 读取表格内容
    if doc.tables:
        print("\n\n表格内容:")
        for i, table in enumerate(doc.tables):
            print(f"\n--- 表格 {i+1} ---")
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                print(" | ".join(row_data))

if __name__ == "__main__":
    # 读取prompt_v2.docx
    read_docx(r"d:\wsp\test\ts_mpp_parser_vscode\prompt_v2.docx")
    
    print("\n\n" + "="*80 + "\n\n")
    
    # 读取TsMpp开发参考.docx
    read_docx(r"d:\wsp\test\ts_mpp_parser_vscode\TsMpp开发参考.docx")
